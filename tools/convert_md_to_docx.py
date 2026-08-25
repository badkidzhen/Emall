import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"D:\111文档\技师学院\项目\商城\分销商城(2).md")
OUTPUT = Path(r"D:\py-project\Emall\分销商城(2).docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    cols = len(table.columns)
    if not cols:
        return
    width = int(9360 / cols)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / cols)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_inline_markdown(paragraph, text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("1F4D78")
        else:
            paragraph.add_run(part)


def parse_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith("|") or "|" not in stripped[1:]:
            break
        table_lines.append(stripped)
        i += 1
    if len(table_lines) < 2 or not re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", table_lines[1]):
        return None, start

    rows = []
    for line in [table_lines[0], *table_lines[2:]]:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, text)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for run in paragraph.runs:
                    run.bold = True


def flush_paragraph(doc, buffer):
    text = " ".join(part.strip() for part in buffer if part.strip())
    buffer.clear()
    if text:
        add_inline_markdown(doc.add_paragraph(), text)


def add_code_block(doc, code_lines):
    if not code_lines:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        if idx != len(code_lines) - 1:
            paragraph.add_run("\n")


def convert():
    text = SOURCE.read_text(encoding="utf-8-sig")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)

    buffer = []
    in_code = False
    code_lines = []
    title_written = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_paragraph(doc, buffer)
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        table_rows, next_i = parse_table(lines, i)
        if table_rows:
            flush_paragraph(doc, buffer)
            add_table(doc, table_rows)
            i = next_i
            continue

        if not stripped:
            flush_paragraph(doc, buffer)
            i += 1
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph(doc, buffer)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph(doc, buffer)
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            if not title_written and level == 1:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(12)
                run = paragraph.add_run(text)
                run.bold = True
                run.font.name = "Calibri"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor.from_string("0B2545")
                title_written = True
            else:
                add_inline_markdown(doc.add_heading("", level=level), text)
            i += 1
            continue

        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            flush_paragraph(doc, buffer)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline_markdown(paragraph, quote.group(1))
            i += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph(doc, buffer)
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline_markdown(paragraph, (bullet or numbered).group(1).strip())
            i += 1
            continue

        buffer.append(stripped)
        i += 1

    flush_paragraph(doc, buffer)
    add_code_block(doc, code_lines)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    convert()
